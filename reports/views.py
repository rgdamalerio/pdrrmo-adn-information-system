import json
import os
import locale
import matplotlib
matplotlib.use('agg') 
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from django.shortcuts import render
from reports.models import *
from django.db.models import Sum
from django.http import HttpResponse, JsonResponse
from django.core import serializers

def export_infant_counts(request):
    infant_counts = FloodReport.objects.values('municipality_name').annotate(
        total_infants=Sum('num_male_infant') + Sum('num_female_infant')
    )

    chart_data = []
    count_infants = 0
    for result in infant_counts:
        municipality_name = result['municipality_name']
        total_infants = result['total_infants']
        count_infants += total_infants
        chart_data.append({'Municipality': municipality_name, 'Total Infants': total_infants})

    # Generate the chart using matplotlib
    labels = [data['Municipality'] for data in chart_data]
    values = [data['Total Infants'] for data in chart_data]
    fig = plt.figure(figsize=(12, 10))
    plt.pie(values, labels=labels, autopct='%1.1f%%', textprops={'fontsize': 12})

    plt.title('Infant Counts by Municipality (Flood Hazard)')
    plt.tight_layout()

    # Set the download path
    download_path = os.path.join(os.path.expanduser('~'), 'Downloads')

    # Save the chart as an image
    chart_image_path = os.path.join(download_path, 'chart.png')
    plt.savefig(chart_image_path)
    plt.close()

    # Define the template file path
    template_file = 'static/template_files/floodreport_template.docx'

    # Create a Word document
    document = Document(template_file)

    # Format the total_infants value with thousands separators
    locale.setlocale(locale.LC_ALL, '')  # Use the system's default locale
    for i, paragraph in enumerate(document.paragraphs):
        if '{{ infant_counts }}' in paragraph.text:
            total_infants_str = locale.format_string('%d', count_infants, grouping=True)
            paragraph.text = paragraph.text.replace('{{ infant_counts }}', total_infants_str)
            document.paragraphs[i].insert_paragraph_before().add_run().add_picture(chart_image_path, width=Inches(7), height=Inches(5))
            break

    # Save the Word document
    document_path = os.path.join(download_path, 'document.docx')
    document.save(document_path)

    # Return the document as an HTTP response
    with open(document_path, 'rb') as file:
        response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=document.docx'
        return response